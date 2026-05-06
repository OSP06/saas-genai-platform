import re
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import structlog
from sqlalchemy import text, select, func, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.exceptions import DocumentNotFoundError, EmbeddingError
from app.models.document import Document, DocumentChunk
from app.schemas.rag import (
    Citation,
    DocumentListItem,
    DocumentStatusResponse,
    DocumentUploadResponse,
)
from app.services.embedding_service import EmbeddingBackend
from app.services.llm_service import LLMService, LLMTokenUsage
from app.services.storage_service import StorageBackend

import asyncio

logger = structlog.get_logger()
settings = get_settings()

# Limit concurrent document ingestion to avoid thread-pool and DB saturation.
_ingestion_semaphore = asyncio.Semaphore(3)

EXTENSION_TO_TYPE = {
    ".pdf": "PDF",
    ".docx": "Word",
    ".doc": "Word",
    ".txt": "Text",
    ".md": "Markdown",
}


def _format_size(size_bytes: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024  # type: ignore[assignment]
    return f"{size_bytes:.1f} TB"


class RAGService:
    def __init__(
        self,
        db: AsyncSession,
        embedding_svc: EmbeddingBackend,
        storage: StorageBackend,
        llm: LLMService,
    ):
        self._db = db
        self._embedding = embedding_svc
        self._storage = storage
        self._llm = llm

    # ------------------------------------------------------------------
    # Document management
    # ------------------------------------------------------------------

    async def save_document_record(
        self, filename: str, size_bytes: int, suffix: str
    ) -> Document:
        """INSERT a Document row with status=pending and return it."""
        doc = Document(
            id=uuid.uuid4(),
            name=filename,
            type=EXTENSION_TO_TYPE.get(suffix.lower(), "Text"),
            size_bytes=size_bytes,
            file_path="",  # will be set in ingest_document
            status="pending",
        )
        self._db.add(doc)
        await self._db.flush()
        logger.info("document_record_created", doc_id=str(doc.id), name=filename)
        return doc

    async def list_documents(self, user_id: str = "default") -> list[DocumentListItem]:
        """Fetch all documents for a user with chunk counts."""
        result = await self._db.execute(
            text(
                """
                SELECT d.id, d.name, d.type, d.size_bytes, d.status,
                       d.created_at, COUNT(c.id) AS chunk_count
                FROM documents d
                LEFT JOIN document_chunks c ON c.document_id = d.id
                WHERE d.user_id = :user_id
                GROUP BY d.id
                ORDER BY d.created_at DESC
                """
            ),
            {"user_id": user_id},
        )
        rows = result.mappings().all()
        items = []
        for row in rows:
            items.append(
                DocumentListItem(
                    id=row["id"],
                    name=row["name"],
                    type=row["type"],
                    size=_format_size(row["size_bytes"]),
                    chunks=int(row["chunk_count"]),
                    uploadedAt=row["created_at"],
                )
            )
        return items

    async def get_document_status(self, document_id: uuid.UUID) -> DocumentStatusResponse:
        result = await self._db.execute(
            select(Document).where(Document.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if doc is None:
            raise DocumentNotFoundError(str(document_id))

        chunk_count_result = await self._db.execute(
            text("SELECT COUNT(*) FROM document_chunks WHERE document_id = :doc_id"),
            {"doc_id": document_id},
        )
        chunk_count = chunk_count_result.scalar() or 0

        progress: float | None = None
        if doc.status == "processing":
            progress = min(0.9, chunk_count / 20.0) if chunk_count else 0.05
        elif doc.status == "ready":
            progress = 1.0

        return DocumentStatusResponse(
            id=doc.id,
            status=doc.status,
            progress=progress,
            chunksCreated=int(chunk_count) if doc.status == "ready" else None,
            error=doc.error_msg,
        )

    async def delete_document(self, document_id: uuid.UUID) -> bool:
        result = await self._db.execute(
            select(Document).where(Document.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if doc is None:
            raise DocumentNotFoundError(str(document_id))

        file_path = doc.file_path
        await self._db.execute(delete(Document).where(Document.id == document_id))
        await self._db.flush()

        if file_path:
            try:
                await self._storage.delete(file_path)
            except Exception as e:
                logger.warning("file_delete_failed", path=file_path, error=str(e))

        logger.info("document_deleted", doc_id=str(document_id))
        return True

    # ------------------------------------------------------------------
    # Ingestion pipeline (runs as a background task)
    # ------------------------------------------------------------------

    async def ingest_document(
        self,
        document_id: uuid.UUID,
        file_bytes: bytes,
        file_type: str,
    ) -> None:
        """
        Full ingestion pipeline:
        1. Save file to storage
        2. Extract text per page
        3. Chunk text with sliding window
        4. Batch embed chunks
        5. Bulk INSERT document_chunks
        6. Mark document ready
        """
        from app.database import AsyncSessionLocal

        async with _ingestion_semaphore:
            async with AsyncSessionLocal() as session:
                try:
                    # Mark processing
                    await session.execute(
                        text("UPDATE documents SET status='processing', updated_at=NOW() WHERE id=:id"),
                        {"id": document_id},
                    )
                    await session.commit()

                    # Save file
                    file_path = await self._storage.save(
                        file_bytes,
                        f"{document_id}{file_type}",
                        "documents",
                    )
                    await session.execute(
                        text("UPDATE documents SET file_path=:path WHERE id=:id"),
                        {"path": file_path, "id": document_id},
                    )
                    await session.commit()

                    # Extract text
                    pages = await self._extract_text(file_bytes, file_type)

                    # Chunk with page tracking
                    all_chunks: list[tuple[str, int | None, int]] = []
                    chunk_idx = 0
                    for page_text, page_number in pages:
                        if not page_text.strip():
                            continue
                        chunks = self._chunk_text(page_text)
                        for chunk_text in chunks:
                            all_chunks.append((chunk_text, page_number, chunk_idx))
                            chunk_idx += 1

                    if not all_chunks:
                        raise ValueError("No text content extracted from document")

                    # Batch embed (32 chunks per batch)
                    batch_size = 32
                    chunk_texts = [c[0] for c in all_chunks]
                    all_embeddings: list[list[float]] = []
                    for i in range(0, len(chunk_texts), batch_size):
                        batch = chunk_texts[i : i + batch_size]
                        embeddings = await self._embedding.embed(batch)
                        all_embeddings.extend(embeddings)

                    # Bulk INSERT in batches of 50 rows to minimize DB round trips
                    chunk_rows = []
                    for (chunk_text, page_number, idx), embedding in zip(all_chunks, all_embeddings):
                        vec_str = "[" + ",".join(str(v) for v in embedding) + "]"
                        chunk_rows.append({
                            "id": str(uuid.uuid4()),
                            "doc_id": str(document_id),
                            "content": chunk_text,
                            "idx": idx,
                            "page": page_number,
                            "embedding": vec_str,
                        })

                    insert_sql = text(
                        """
                        INSERT INTO document_chunks
                            (id, document_id, content, chunk_index, page_number, embedding, created_at)
                        VALUES
                            (:id, :doc_id, :content, :idx, :page, :embedding::vector, NOW())
                        """
                    )
                    for i in range(0, len(chunk_rows), 50):
                        await session.execute(insert_sql, chunk_rows[i : i + 50])

                    await session.execute(
                        text(
                            "UPDATE documents SET status='ready', updated_at=NOW() WHERE id=:id"
                        ),
                        {"id": document_id},
                    )
                    await session.commit()
                    logger.info(
                        "document_ingested",
                        doc_id=str(document_id),
                        chunks=len(all_chunks),
                    )

                except Exception as exc:
                    await session.rollback()
                    await session.execute(
                        text(
                            "UPDATE documents SET status='error', error_msg=:err, updated_at=NOW() WHERE id=:id"
                        ),
                        {"err": str(exc)[:500], "id": document_id},
                    )
                    await session.commit()
                    logger.error("document_ingest_failed", doc_id=str(document_id), error=str(exc))

    async def _extract_text(
        self, file_bytes: bytes, file_type: str
    ) -> list[tuple[str, int | None]]:
        """Returns [(page_text, page_number | None), ...]."""
        import asyncio

        loop = asyncio.get_running_loop()

        if file_type in (".pdf",):
            return await loop.run_in_executor(None, self._extract_pdf, file_bytes)
        elif file_type in (".docx", ".doc"):
            return await loop.run_in_executor(None, self._extract_docx, file_bytes)
        else:
            return [(file_bytes.decode("utf-8", errors="replace"), None)]

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> list[tuple[str, int]]:
        import io
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append((text, i))
        return pages

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> list[tuple[str, None]]:
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [(text, None)] if text else []

    def _chunk_text(self, text: str) -> list[str]:
        """
        Sliding-window chunker on word boundaries.
        Tries to split on double-newlines first (paragraph breaks),
        then falls back to char-level sliding window.
        """
        chunk_size = settings.CHUNK_SIZE
        overlap = settings.CHUNK_OVERLAP

        # Split on paragraph breaks first
        paragraphs = re.split(r"\n\n+", text.strip())
        chunks: list[str] = []
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 1 <= chunk_size:
                current = (current + "\n\n" + para).strip()
            else:
                if current:
                    chunks.append(current)
                # If single paragraph is too long, slice it
                if len(para) > chunk_size:
                    for start in range(0, len(para), chunk_size - overlap):
                        piece = para[start : start + chunk_size]
                        if piece.strip():
                            chunks.append(piece.strip())
                    current = ""
                else:
                    current = para

        if current:
            chunks.append(current)

        return chunks or [text[:chunk_size]]

    # ------------------------------------------------------------------
    # Query (vector search + grounded answer)
    # ------------------------------------------------------------------

    async def query(
        self,
        query_text: str,
        document_ids: list[uuid.UUID] | None,
        max_citations: int,
    ) -> tuple[str, list[Citation], LLMTokenUsage | None]:
        """
        1. Embed query
        2. Cosine similarity search in pgvector
        3. Build grounded prompt
        4. LLM completion
        5. Return (answer, citations)
        """
        _query_start = time.perf_counter()
        try:
            query_embedding = (await self._embedding.embed([query_text]))[0]
        except Exception as e:
            raise EmbeddingError(f"Query embedding failed: {e}") from e

        vec_str = "[" + ",".join(str(v) for v in query_embedding) + "]"

        doc_filter = ""
        params: dict[str, Any] = {
            "query_vec": vec_str,
            "threshold": settings.SIMILARITY_THRESHOLD,
            "limit": max_citations,
        }
        if document_ids:
            doc_filter = "AND dc.document_id = ANY(:doc_ids)"
            params["doc_ids"] = [str(d) for d in document_ids]

        sql = text(
            f"""
            SELECT
                dc.id,
                dc.document_id,
                dc.content,
                dc.page_number,
                d.name AS document_name,
                1 - (dc.embedding <=> :query_vec::vector) AS score
            FROM document_chunks dc
            JOIN documents d ON d.id = dc.document_id
            WHERE d.status = 'ready'
              AND dc.embedding IS NOT NULL
              {doc_filter}
            ORDER BY dc.embedding <=> :query_vec::vector
            LIMIT :limit
            """
        )

        # Boost recall at scale: probe more IVFFlat clusters before searching
        await self._db.execute(text("SET LOCAL ivfflat.probes = 10"))
        result = await self._db.execute(sql, params)
        rows = result.mappings().all()

        if not rows:
            return (
                "I couldn't find any relevant information in the knowledge base to answer your question.",
                [],
                None,
            )

        # Keep only chunks that meet the similarity threshold.
        # If nothing meets the bar, tell the user rather than returning low-quality results.
        filtered = [r for r in rows if float(r["score"]) >= settings.SIMILARITY_THRESHOLD]
        if not filtered:
            return (
                "No sufficiently relevant content was found in the knowledge base. "
                "Try rephrasing your question or uploading more relevant documents.",
                [],
                None,
            )

        system_prompt, user_prompt = self._build_rag_prompt(query_text, filtered)
        answer, usage = await self._llm.complete(
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
            temperature=0.3,
        )

        citations = [
            Citation(
                id=row["id"],
                documentId=row["document_id"],
                documentName=row["document_name"],
                content=row["content"][:500],
                page=row["page_number"],
                relevanceScore=round(float(row["score"]), 4),
            )
            for row in filtered
        ]

        logger.info(
            "rag_query_complete",
            query_len=len(query_text),
            citations=len(citations),
            model=usage.model,
            tokens=usage.input_tokens + usage.output_tokens,
            latency_ms=int((time.perf_counter() - _query_start) * 1000),
        )
        return answer, citations, usage

    def _build_rag_prompt(
        self, query: str, chunks: list[Any]
    ) -> tuple[str, str]:
        system = (
            "You are a precise, grounded Q&A assistant. "
            "Answer the user's question using ONLY the provided context excerpts. "
            "If the context does not contain sufficient information, say so clearly. "
            "Cite sources inline as [Doc: <document_name>, p.<page>] when page is available, "
            "or [Doc: <document_name>] otherwise."
        )

        context_parts = []
        for i, chunk in enumerate(chunks, start=1):
            page_info = f", page {chunk['page_number']}" if chunk["page_number"] else ""
            context_parts.append(
                f"[{i}] From '{chunk['document_name']}'{page_info}:\n{chunk['content']}"
            )

        user = f"Context:\n\n" + "\n\n---\n\n".join(context_parts) + f"\n\nQuestion: {query}"
        return system, user
